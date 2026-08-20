"""
Document Generator for Week 2 Task: EDA and Visualization Framework Design.
Generates:
1. EDA_and_Visualization_Framework.docx (Native Word Document with full styling, strictly < 2048 KB)
2. EDA_and_Visualization_Framework.doc (Rich HTML-formatted Document for Word/Web)
"""

import os
import base64
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "output_visuals")

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_callout_box(doc, text_list, title="KEY TAKEAWAY", border_color="1B365D", bg_color="F4F6F9"):
    """Adds a stylish callout box to the docx."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    
    cell = table.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"📌 {title}\n")
    run_title.font.bold = True
    run_title.font.size = Pt(10.5)
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    for item in text_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_before = Pt(2)
        p_item.paragraph_format.space_after = Pt(2)
        r = p_item.add_run(f"• {item}")
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    doc.add_paragraph()

def add_code_snippet(doc, code_str, language="Python"):
    """Adds a formatted code block to docx."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(6.5)
    
    cell = table.cell(0, 0)
    set_cell_background(cell, "282C34")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    lang_run = p.add_run(f"# [{language} Implementation]\n")
    lang_run.font.name = "Consolas"
    lang_run.font.size = Pt(8.5)
    lang_run.font.color.rgb = RGBColor(0x98, 0xC3, 0x79)
    
    code_run = p.add_run(code_str)
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(8.5)
    code_run.font.color.rgb = RGBColor(0xAB, 0xB2, 0xBF)
    
    doc.add_paragraph()

def format_table(table, header_bg="1B365D", alt_bg="F8F9FA"):
    """Applies executive styling to docx tables."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        set_cell_background(cell, header_bg)
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9.5)
    
    for r_idx, row in enumerate(table.rows[1:], start=1):
        bg = alt_bg if r_idx % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

def generate_docx():
    print("[*] Generating EDA_and_Visualization_Framework.docx...")
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # Document Header / Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_sub = title_p.add_run("DATA SCIENCE TECHNICAL BLUEPRINT • WEEK 2 DELIVERABLE\n")
    run_sub.font.bold = True
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = RGBColor(0xE0, 0x5A, 0x47)
    
    run_title = title_p.add_run("Exploratory Data Analysis (EDA) & Visualization Framework Design")
    run_title.font.bold = True
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(16)
    meta_run = meta_p.add_run("A Universal, Multi-Phase Analytical Methodology for Tabular, Structured, and Complex Datasets\nAuthor: Data Science & AI Research Team | Status: Production Blueprint")
    meta_run.font.italic = True
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(14)
    r_div = p_div.add_run("━" * 55)
    r_div.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)

    # SECTION 1
    h1 = doc.add_heading("1. Introduction to Exploratory Data Analysis (EDA)", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    p = doc.add_paragraph(
        "Exploratory Data Analysis (EDA) is the foundational, non-negotiable first phase of any rigorous Data Science "
        "and Machine Learning lifecycle. Pioneered by statistician John Tukey in 1977, EDA is an investigative philosophy "
        "and analytical discipline that emphasizes understanding the underlying structure, distributional properties, "
        "anomalies, and relational dynamics of a dataset before imposing parametric assumptions or building predictive models."
    )
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph(
        "In modern enterprise data environments, rushing into feature engineering or model training without exhaustive EDA "
        "frequently leads to severe downstream failure modes, including 'Garbage-In, Garbage-Out' (GIGO) phenomena, silent data "
        "leakage, multicollinearity degradation, bias amplification, and catastrophic model drift in production. "
        "EDA transforms raw, unvetted data into an audited, structured, and statistically grounded asset."
    )
    p.paragraph_format.space_after = Pt(10)

    obj_table = doc.add_table(rows=6, cols=3)
    obj_table.columns[0].width = Inches(1.5)
    obj_table.columns[1].width = Inches(2.2)
    obj_table.columns[2].width = Inches(2.8)
    
    headers = ["Core Dimension", "Primary Objective", "Business & Engineering Impact"]
    for i, h in enumerate(headers):
        obj_table.cell(0, i).paragraphs[0].text = h
        
    rows_data = [
        ("Structural Auditing", "Assess dimensions, feature types, memory footprint, and data schemas.", "Prevents runtime schema errors and optimizes memory allocation."),
        ("Data Quality Assessment", "Identify missingness patterns (MCAR/MAR/MNAR), anomalies, and corruption.", "Determines valid imputation strategies and eliminates corrupted records."),
        ("Distributional Profiling", "Quantify central tendency, dispersion, skewness, and modality.", "Informs necessary non-linear transformations (Log, Box-Cox, Scaling)."),
        ("Hypothesis & Pattern Mining", "Uncover non-obvious correlations, multi-way interactions, and clusters.", "Guides high-yield domain feature engineering and hypothesis testing."),
        ("Assumptions Validation", "Test statistical assumptions (normality, homoscedasticity, linearity).", "Ensures correct model selection (Parametric vs. Non-Parametric).")
    ]
    for r_idx, row in enumerate(rows_data, start=1):
        for c_idx, val in enumerate(row):
            obj_table.cell(r_idx, c_idx).paragraphs[0].text = val
            
    format_table(obj_table)
    doc.add_paragraph()

    add_callout_box(
        doc,
        [
            "EDA is iterative, non-linear, and empirical: Findings in bivariate analysis often demand returning to data cleaning.",
            "Visual diagnostics must always be paired with quantitative summary statistics (e.g., skewness, p-values).",
            "A well-documented EDA serves as the technical single-source-of-truth for both engineers and business stakeholders."
        ],
        title="CORE PRINCIPLES OF ROBUST EDA"
    )

    # SECTION 2
    h1 = doc.add_heading("2. Data Taxonomy & Type-Specific Exploration Strategies", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    p = doc.add_paragraph(
        "A universal EDA framework must establish a formal taxonomy of data types. Analytical methods, statistical tests, "
        "and chart types are strictly dictated by the mathematical properties of each variable category:"
    )
    p.paragraph_format.space_after = Pt(8)

    dtype_table = doc.add_table(rows=6, cols=4)
    dtype_table.columns[0].width = Inches(1.3)
    dtype_table.columns[1].width = Inches(1.5)
    dtype_table.columns[2].width = Inches(1.8)
    dtype_table.columns[3].width = Inches(1.9)
    
    dt_headers = ["Data Type", "Definition & Properties", "Mathematical Operations", "Recommended Visualizations"]
    for i, h in enumerate(dt_headers):
        dtype_table.cell(0, i).paragraphs[0].text = h
        
    dt_rows = [
        ("Nominal (Categorical)", "Unordered discrete classes (e.g., Gender, State, Device Type).", "Equality (==, !=), Mode, Frequency Counts.", "Bar charts, Donut charts, Count plots."),
        ("Ordinal (Categorical)", "Ordered discrete categories with ranking (e.g., Education, Rating).", "Comparison (<, >), Median, Percentiles, Rank correlation.", "Ordered Bar charts, Spine plots, Heatmaps."),
        ("Discrete Numeric", "Countable integer quantities (e.g., Number of Dependents, Visits).", "Mean, Median, Standard Deviation, Poisson modeling.", "Step Histograms, Count plots, Stem-and-leaf."),
        ("Continuous Numeric", "Unbroken scale of measurements (e.g., Revenue, Tenure, Speed).", "Full arithmetic, Variance, Kurtosis, Calculus transformations.", "Histograms + KDE, Box plots, ECDF, Violin plots."),
        ("Temporal / Datetime", "Time-series timestamps, cycles (e.g., Transaction Date, Hour).", "Lags, Differences, Seasonality decomposition, Fourier.", "Line charts, Horizon plots, Calendar heatmaps.")
    ]
    for r_idx, row in enumerate(dt_rows, start=1):
        for c_idx, val in enumerate(row):
            dtype_table.cell(r_idx, c_idx).paragraphs[0].text = val
            
    format_table(dtype_table)
    doc.add_paragraph()

    # SECTION 3
    h1 = doc.add_heading("3. Data Quality, Missing Value Diagnostics & Outlier Detection", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    p = doc.add_paragraph(
        "Data hygiene determines the integrity of all downstream conclusions. Real-world enterprise datasets invariably "
        "contain missing values, data entry anomalies, corrupted formats, and extreme outliers."
    )
    
    h2 = doc.add_heading("3.1 Theoretical Taxonomy of Missing Data Mechanisms", level=2)
    for r in h2.runs:
        r.font.color.rgb = RGBColor(0x4B, 0x6B, 0x94)
        
    p1 = doc.add_paragraph(
        "1. Missing Completely at Random (MCAR): The probability of missingness is completely independent of both observed "
        "and unobserved variables. (e.g., dropped network packet). Safe for listwise deletion or simple mean/median imputation.\n"
        "2. Missing at Random (MAR): The probability of missingness depends systematically on observed data, but not on the "
        "missing value itself. Requires conditional imputation (KNN, MICE, IterativeImputer).\n"
        "3. Missing Not at Random (MNAR): Missingness depends directly on the unobserved value itself. Requires missingness indicator flags."
    )
    p1.paragraph_format.space_after = Pt(8)

    h2 = doc.add_heading("3.2 Outlier Detection Methodologies", level=2)
    for r in h2.runs:
        r.font.color.rgb = RGBColor(0x4B, 0x6B, 0x94)

    outlier_code = (
        "# 1. Tukey's Interquartile Range (IQR) Rule\n"
        "Q1 = df['feature'].quantile(0.25)\n"
        "Q3 = df['feature'].quantile(0.75)\n"
        "IQR = Q3 - Q1\n"
        "lower_bound = Q1 - 1.5 * IQR\n"
        "upper_bound = Q3 + 1.5 * IQR\n"
        "outliers_iqr = df[(df['feature'] < lower_bound) | (df['feature'] > upper_bound)]\n\n"
        "# 2. Standardized Z-Score Method (Parametric / Gaussian)\n"
        "z_scores = np.abs((df['feature'] - df['feature'].mean()) / df['feature'].std())\n"
        "outliers_z = df[z_scores > 3.0]\n\n"
        "# 3. Machine Learning Isolation Forest (Multivariate Non-Parametric)\n"
        "from sklearn.ensemble import IsolationForest\n"
        "iso = IsolationForest(contamination=0.02, random_state=42)\n"
        "df['anomaly_flag'] = iso.fit_predict(df[numeric_cols]) # -1 = Outlier"
    )
    add_code_snippet(doc, outlier_code, language="Python - Outlier Detection Protocols")

    # SECTION 4
    h1 = doc.add_heading("4. The Multi-Tier Systematic Analysis Framework", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    h2 = doc.add_heading("4.1 Tier 1: Univariate Analysis (Single Variable Profiling)", level=2)
    for r in h2.runs:
        r.font.color.rgb = RGBColor(0x4B, 0x6B, 0x94)
    p = doc.add_paragraph(
        "Examines variables in isolation. For continuous variables, metrics include Mean, Median, Variance, Skewness "
        "(measure of asymmetry: Skew > 1 indicates right-tail skew), and Kurtosis (measure of heavy tails). "
        "For categorical variables, assess frequency counts, relative proportions, and cardinality."
    )

    h2 = doc.add_heading("4.2 Tier 2: Bivariate Analysis (Pairwise Relationships)", level=2)
    for r in h2.runs:
        r.font.color.rgb = RGBColor(0x4B, 0x6B, 0x94)
    p = doc.add_paragraph(
        "Explores interactions between pairs of variables:\n"
        "• Continuous vs. Continuous: Pearson Correlation, Spearman Rank Correlation, Scatter plots with regression trendlines.\n"
        "• Categorical vs. Continuous: Grouped Box/Violin plots, Independent Welch's t-test (2 groups), One-Way ANOVA / Kruskal-Wallis (>2 groups).\n"
        "• Categorical vs. Categorical: Cross-tabulations, Stacked Percentage Bar charts, Pearson's Chi-Square Test (χ²), and Cramer's V."
    )

    h2 = doc.add_heading("4.3 Tier 3: Multivariate Analysis & Interaction Dynamics", level=2)
    for r in h2.runs:
        r.font.color.rgb = RGBColor(0x4B, 0x6B, 0x94)
    p = doc.add_paragraph(
        "Investigates simultaneous interactions across three or more features:\n"
        "• Correlation Heatmaps: Masked lower-triangle heatmaps to detect multicollinearity (VIF > 5.0 indicates redundant features).\n"
        "• Faceted Multi-Condition Scatter Plots: Plotting X vs. Y with Hue (color), Size, and Style encodings.\n"
        "• Dimensionality Reduction Projections: Principal Component Analysis (PCA) or t-SNE / UMAP for high-dimensional clustering."
    )

    # SECTION 5
    h1 = doc.add_heading("5. Visualization Taxonomy & Plot Selection Decision Matrix", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    viz_table = doc.add_table(rows=8, cols=4)
    viz_table.columns[0].width = Inches(1.3)
    viz_table.columns[1].width = Inches(1.4)
    viz_table.columns[2].width = Inches(1.8)
    viz_table.columns[3].width = Inches(2.0)

    v_headers = ["Analytical Goal", "Variable Configuration", "Recommended Visual Chart", "Best Practice & Insight Generated"]
    for i, h in enumerate(v_headers):
        viz_table.cell(0, i).paragraphs[0].text = h

    v_rows = [
        ("Distribution Shape", "1 Continuous", "Histogram + KDE Overlay", "Set appropriate bin widths; overlay mean/median reference lines."),
        ("Spread & Outliers", "1 Continuous", "Box Plot / Violin Plot", "Visualizes IQR, 1.5x whisker boundaries, and distribution density."),
        ("Category Frequency", "1 Categorical", "Horizontal / Vertical Bar Chart", "Sort bars by descending frequency; annotate exact % labels."),
        ("Continuous Association", "2 Continuous", "Scatter Plot + Lowess Trend", "Identifies linear/non-linear correlation, clustering, and heteroscedasticity."),
        ("Grouped Comparison", "1 Cat + 1 Continuous", "Split Violin / Box Plot", "Compares median shift, spread, and tail shape across categories."),
        ("Categorical Interaction", "2 Categorical", "Stacked 100% Bar / Mosaic", "Directly compares relative composition and contingent probabilities."),
        ("Multivariate Correlation", "3+ Continuous", "Masked Heatmap with Annotations", "Use diverging colormaps (coolwarm) with values restricted to [-1, 1].")
    ]
    for r_idx, row in enumerate(v_rows, start=1):
        for c_idx, val in enumerate(row):
            viz_table.cell(r_idx, c_idx).paragraphs[0].text = val

    format_table(viz_table)
    doc.add_paragraph()

    # SECTION 6
    h1 = doc.add_heading("6. Python Tooling & Library Ecosystem Deep Dive", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    tools_table = doc.add_table(rows=8, cols=3)
    tools_table.columns[0].width = Inches(1.3)
    tools_table.columns[1].width = Inches(2.3)
    tools_table.columns[2].width = Inches(2.9)

    t_headers = ["Library / Tool", "Core Architectural Role", "Key Functions & Unique Capabilities"]
    for i, h in enumerate(t_headers):
        tools_table.cell(0, i).paragraphs[0].text = h

    t_rows = [
        ("Pandas (v2.x+)", "Data manipulation, vectorized transformations, and structural profiling.", "read_csv(), describe(), info(), groupby(), pivot_table(), value_counts(), qcut()."),
        ("NumPy (v2.x+)", "High-performance N-dimensional array mathematics.", "Vectorized boolean masking, quantile(), log1p(), where(), corrcoef()."),
        ("Matplotlib", "Low-level imperative visualization engine and figure architecture.", "Custom figure grids (subplots), custom axis formatting, fine-grained DPI and layout control."),
        ("Seaborn", "High-level statistical visualization library built on Matplotlib.", "Automated statistical aggregation: histplot(), kdeplot(), violinplot(), heatmap(), pairplot()."),
        ("Plotly / Dash", "Declarative, interactive WebGL-rendered charts.", "Hover tooltips, dynamic zooming/panning, 3D surface plots, interactive HTML dashboards."),
        ("SciPy (scipy.stats)", "Rigorous inferential statistics and hypothesis testing.", "ttest_ind(), chi2_contingency(), f_oneway(), shapiro(), normaltest(), stats.boxcox()."),
        ("Missingno", "Dedicated visual diagnostics for missing data matrices.", "matrix(), bar(), heatmap(), dendrogram() to inspect MCAR/MAR patterns.")
    ]
    for r_idx, row in enumerate(t_rows, start=1):
        for c_idx, val in enumerate(row):
            tools_table.cell(r_idx, c_idx).paragraphs[0].text = val

    format_table(tools_table)
    doc.add_paragraph()

    # SECTION 7: CASE STUDY
    h1 = doc.add_heading("7. Practical End-to-End Demonstration: Telco Churn Dataset", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    p = doc.add_paragraph(
        "To illustrate the end-to-end framework in action, we executed our automated EDA pipeline on the IBM Telco Customer Churn dataset, "
        "downloaded directly via terminal. The dataset consists of 7,043 customer records and 21 heterogeneous features covering customer demographics, "
        "subscribed services, billing metrics, and contract duration."
    )
    p.paragraph_format.space_after = Pt(8)

    # Insert Fig 1
    fig1_path = os.path.join(OUTPUT_DIR, "fig1_missing_and_target.png")
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(5.8))
        cap = doc.add_paragraph("Figure 1: Missing Data Diagnostics & Target Class Balance (26.5% Churn Rate).")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True

    p = doc.add_paragraph(
        "Analytical Findings from Phase 1 & 2:\n"
        "• Data Hygiene Issue Detected: 'TotalCharges' was stored as an Object dtype due to 11 records containing whitespace (' '). "
        "Converting to numeric revealed these were brand-new customers with tenure = 0 months. Imputed with the median ($1,397.48).\n"
        "• Class Imbalance: Churn rate is 26.5% (1,869 churners vs 5,174 retained). Accuracy alone will be misleading; Precision-Recall AUC and ROC-AUC are required."
    )

    # Insert Fig 2
    fig2_path = os.path.join(OUTPUT_DIR, "fig2_univariate_numeric.png")
    if os.path.exists(fig2_path):
        doc.add_picture(fig2_path, width=Inches(5.8))
        cap = doc.add_paragraph("Figure 2: Univariate Distribution Analysis for Tenure, Monthly Charges, and Total Charges.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True

    p = doc.add_paragraph(
        "Numerical Distribution Profiling:\n"
        "• Tenure: Exhibits a bimodal distribution with massive peaks at Month 1 (new acquisitions) and Month 72 (long-term loyalists). Mean: 32.37m, Median: 29.0m, Skewness: 0.24.\n"
        "• Monthly Charges: Displays a large cluster around $20 and a broad high-tier distribution between $70-$105. Mean: $64.76, Median: $70.35.\n"
        "• Total Charges: Strong right-skewed distribution (Skew: +0.96) reflecting compounding tenure and bill amounts."
    )

    # Insert Fig 3
    fig3_path = os.path.join(OUTPUT_DIR, "fig3_univariate_categorical.png")
    if os.path.exists(fig3_path):
        doc.add_picture(fig3_path, width=Inches(5.8))
        cap = doc.add_paragraph("Figure 3: Univariate Proportions for High-Impact Categorical Features.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True

    p = doc.add_paragraph(
        "Categorical Feature Insights:\n"
        "• Contract Structure: 55.0% Month-to-Month contracts, 24.1% 2-Year contracts, and 20.9% 1-Year contracts.\n"
        "• Internet Service: 44.0% Fiber Optic, 34.4% DSL, and 21.6% no internet service.\n"
        "• Payment Methods: Electronic Check is the most frequent payment method (33.6%)."
    )

    # Insert Fig 4
    fig4_path = os.path.join(OUTPUT_DIR, "fig4_bivariate_continuous_target.png")
    if os.path.exists(fig4_path):
        doc.add_picture(fig4_path, width=Inches(5.8))
        cap = doc.add_paragraph("Figure 4: Bivariate Violin Plots: Continuous Feature Distributions Segmented by Churn Status.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True

    p = doc.add_paragraph(
        "Bivariate & Inferential Hypothesis Testing Results:\n"
        "• Welch's Independent T-Test (Tenure across Churn): t = -34.82, p-value = 1.20 × 10⁻²³². Churners have an overwhelmingly lower median tenure (10 months) compared to retained customers (38 months).\n"
        "• Monthly Charges Effect: Churners exhibit significantly higher median monthly charges ($79.65 vs. $64.40), indicating price sensitivity is a major churn driver."
    )

    # Insert Fig 5
    fig5_path = os.path.join(OUTPUT_DIR, "fig5_correlation_heatmap.png")
    if os.path.exists(fig5_path):
        doc.add_picture(fig5_path, width=Inches(5.2))
        cap = doc.add_paragraph("Figure 5: Correlation Heatmap across Continuous and Encoded Binary Features.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True

    # Insert Fig 6
    fig6_path = os.path.join(OUTPUT_DIR, "fig6_categorical_churn_crosstabs.png")
    if os.path.exists(fig6_path):
        doc.add_picture(fig6_path, width=Inches(5.8))
        cap = doc.add_paragraph("Figure 6: Cross-Tabulated Churn Rates across Contract Types and Internet Service Tiers.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True

    p = doc.add_paragraph(
        "Cross-Tabulation & Chi-Square Independence Findings:\n"
        "• Chi-Square Test (Contract vs Churn): χ² = 1184.60, p-value = 5.86 × 10⁻²⁵⁸. Month-to-Month contracts have an alarming 42.7% churn rate, compared to just 11.3% for One-Year and 2.8% for Two-Year contracts.\n"
        "• Internet Service Risk: Fiber Optic users churn at 41.9%, compared to 18.9% for DSL and 7.4% for customers with no internet."
    )

    # Insert Fig 7
    fig7_path = os.path.join(OUTPUT_DIR, "fig7_multivariate_interaction.png")
    if os.path.exists(fig7_path):
        doc.add_picture(fig7_path, width=Inches(5.8))
        cap = doc.add_paragraph("Figure 7: Multivariate Interaction Plot (Tenure vs Monthly Charges faceted by Contract & Churn).")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True

    add_callout_box(
        doc,
        [
            "Feature Engineering Opportunity 1: Create 'Tenure_Cohort' (<12m, 12-24m, 24-48m, 48m+) to capture the non-linear risk decay.",
            "Feature Engineering Opportunity 2: Create 'Charge_Ratio' = (TotalCharges / (tenure * MonthlyCharges)) to detect billing anomalies.",
            "Feature Engineering Opportunity 3: Create 'HighRisk_Flag' = (Month-to-Month Contract + Electronic Check + Fiber Optic)."
        ],
        title="ACTIONABLE FEATURE ENGINEERING RECOMMENDATIONS DERIVED FROM EDA"
    )

    # SECTION 8
    h1 = doc.add_heading("8. Plan for Documenting, Reporting & Stakeholder Presentation", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    rep_table = doc.add_table(rows=5, cols=3)
    rep_table.columns[0].width = Inches(1.5)
    rep_table.columns[1].width = Inches(2.2)
    rep_table.columns[2].width = Inches(2.8)

    r_headers = ["Report Section", "Target Audience & Focus", "Key Deliverables & Artifacts"]
    for i, h in enumerate(r_headers):
        rep_table.cell(0, i).paragraphs[0].text = h

    r_rows = [
        ("Executive Summary", "C-Suite & Business Leadership: High-level financial impact and strategic recommendations.", "1-page dashboard, churn drivers ranked by business cost, key ROI projections."),
        ("Data Quality & Governance", "Data Engineering & Compliance: Audit trail of missingness, corruption, and schemas.", "Missing data audit tables, schema validation logs, data pipeline patch recommendations."),
        ("Deep Statistical Analysis", "Data Science & ML Engineers: Distribution parameters, correlation matrices, test stats.", "Full Jupyter notebooks, p-value tables, VIF multicollinearity reports, transformation curves."),
        ("Production Roadmap", "Product Managers & Operations: Concrete operational changes and next sprint milestones.", "Feature engineering dictionary, model baseline criteria, tracking metrics (KPI dashboards).")
    ]
    for r_idx, row in enumerate(r_rows, start=1):
        for c_idx, val in enumerate(row):
            rep_table.cell(r_idx, c_idx).paragraphs[0].text = val

    format_table(rep_table)
    doc.add_paragraph()

    add_callout_box(
        doc,
        [
            "Standardized EDA Checklists ensure reproducibility across multi-engineer data science teams.",
            "Automated reporting (e.g. Quarto, Jupyter Book, Streamlit) bridges the gap between raw data and executive strategy.",
            "Always link statistical patterns directly to dollarized business outcomes (e.g., Fiber Optic churn costs $X million annually)."
        ],
        title="FINAL METHODOLOGY SUMMARY & BEST PRACTICES"
    )

    output_docx = os.path.join(SCRIPT_DIR, "EDA_and_Visualization_Framework.docx")
    doc.save(output_docx)
    size_kb = os.path.getsize(output_docx) / 1024
    print(f"[+] Successfully generated '{output_docx}' ({size_kb:.2f} KB)")
    return output_docx

def generate_html_doc():
    """Generates an HTML-based .doc file with rich styling and embedded images."""
    print("[*] Generating EDA_and_Visualization_Framework.doc (HTML-DOC)...")
    
    def get_img_tag(path, caption):
        if not os.path.exists(path):
            return f"<p><em>[Image {path} not found]</em></p>"
        with open(path, "rb") as f:
            encoded = base64.b64encode(f.read()).decode("utf-8")
        return f"""
        <div style="text-align: center; margin: 20px 0;">
            <img src="data:image/png;base64,{encoded}" style="max-width: 95%; height: auto; border: 1px solid #E0E4E8; border-radius: 4px;" />
            <p style="font-size: 9pt; color: #5A6B7C; font-style: italic; margin-top: 5px;">{caption}</p>
        </div>
        """

    fig1_html = get_img_tag(os.path.join(OUTPUT_DIR, "fig1_missing_and_target.png"), "Figure 1: Missing Value Diagnostics & Target Class Balance (26.5% Churn).")
    fig2_html = get_img_tag(os.path.join(OUTPUT_DIR, "fig2_univariate_numeric.png"), "Figure 2: Univariate Numerical Distributions (Histograms + KDE + Boxplot).")
    fig3_html = get_img_tag(os.path.join(OUTPUT_DIR, "fig3_univariate_categorical.png"), "Figure 3: Univariate Proportions for High-Impact Categorical Features.")
    fig4_html = get_img_tag(os.path.join(OUTPUT_DIR, "fig4_bivariate_continuous_target.png"), "Figure 4: Bivariate Violin Plots: Continuous Metrics vs. Churn Target.")
    fig5_html = get_img_tag(os.path.join(OUTPUT_DIR, "fig5_correlation_heatmap.png"), "Figure 5: Correlation Matrix Heatmap across Encoded Binary and Continuous Features.")
    fig6_html = get_img_tag(os.path.join(OUTPUT_DIR, "fig6_categorical_churn_crosstabs.png"), "Figure 6: Cross-Tabulated Churn Proportions across Contract and Internet Tiers.")
    fig7_html = get_img_tag(os.path.join(OUTPUT_DIR, "fig7_multivariate_interaction.png"), "Figure 7: Multivariate Interaction (Tenure vs Monthly Charges by Contract & Churn).")

    html_content = f"""<!DOCTYPE html>
<html xmlns:o="urn:schemas-microsoft-com:office:office"
      xmlns:w="urn:schemas-microsoft-com:office:word"
      xmlns="http://www.w3.org/TR/REC-html40">
<head>
<meta charset="utf-8">
<title>Exploratory Data Analysis (EDA) and Visualization Framework Design</title>
<style>
    @page {{
        size: 8.5in 11in;
        margin: 1.0in 1.0in 1.0in 1.0in;
    }}
    body {{
        font-family: 'Calibri', 'Segoe UI', Arial, sans-serif;
        font-size: 11pt;
        line-height: 1.5;
        color: #2C3E50;
        background-color: #FFFFFF;
        margin: 0;
        padding: 20px;
    }}
    .header-banner {{
        border-bottom: 2px solid #1B365D;
        padding-bottom: 12px;
        margin-bottom: 25px;
    }}
    .tagline {{
        font-size: 10pt;
        font-weight: bold;
        color: #E05A47;
        text-transform: uppercase;
        letter-spacing: 1px;
        margin-bottom: 4px;
    }}
    h1.doc-title {{
        font-size: 22pt;
        color: #1B365D;
        margin: 0 0 6px 0;
        font-weight: 700;
        line-height: 1.2;
    }}
    .doc-meta {{
        font-size: 10pt;
        color: #7F8C8D;
        font-style: italic;
    }}
    h1 {{
        font-size: 15pt;
        color: #1B365D;
        border-bottom: 1.5px solid #E0E4E8;
        padding-bottom: 4px;
        margin-top: 28px;
        margin-bottom: 12px;
    }}
    h2 {{
        font-size: 12.5pt;
        color: #4B6B94;
        margin-top: 18px;
        margin-bottom: 8px;
    }}
    p {{
        margin-top: 0;
        margin-bottom: 10px;
        text-align: justify;
    }}
    .callout {{
        background-color: #F4F6F9;
        border-left: 4px solid #1B365D;
        padding: 12px 16px;
        margin: 16px 0;
        border-radius: 0 4px 4px 0;
    }}
    .callout-title {{
        font-weight: bold;
        font-size: 10.5pt;
        color: #1B365D;
        margin-bottom: 6px;
    }}
    .callout ul {{
        margin: 0;
        padding-left: 18px;
    }}
    .callout li {{
        margin-bottom: 4px;
        font-size: 10pt;
    }}
    table {{
        width: 100%;
        border-collapse: collapse;
        margin: 16px 0;
        font-size: 9.5pt;
    }}
    th {{
        background-color: #1B365D;
        color: #FFFFFF;
        font-weight: 600;
        text-align: left;
        padding: 8px 10px;
        border: 1px solid #1B365D;
    }}
    td {{
        padding: 7px 10px;
        border: 1px solid #E0E4E8;
        vertical-align: top;
    }}
    tr:nth-child(even) {{
        background-color: #F8F9FA;
    }}
    pre.code-block {{
        background-color: #282C34;
        color: #ABB2BF;
        padding: 12px 14px;
        border-radius: 4px;
        font-family: 'Consolas', monospace;
        font-size: 9pt;
        line-height: 1.4;
        overflow-x: auto;
        margin: 14px 0;
    }}
    .code-comment {{ color: #98C379; font-weight: bold; }}
    .code-keyword {{ color: #C678DD; }}
</style>
</head>
<body>

<div class="header-banner">
    <div class="tagline">Data Science Technical Blueprint • Week 2 Deliverable</div>
    <h1 class="doc-title">Exploratory Data Analysis (EDA) & Visualization Framework Design</h1>
    <div class="doc-meta">A Universal, Multi-Phase Analytical Methodology for Tabular, Structured, and Complex Datasets<br>Author: Data Science & AI Research Team</div>
</div>

<h1>1. Executive Summary & Introduction to EDA</h1>
<p>
    <strong>Exploratory Data Analysis (EDA)</strong> is the foundational, non-negotiable first phase of any rigorous Data Science and Machine Learning lifecycle. Pioneered by statistician John Tukey in 1977, EDA represents an investigative philosophy and systematic analytical discipline that emphasizes uncovering the underlying structure, distributional properties, anomalies, and relational dynamics of a dataset before imposing parametric assumptions or building predictive models.
</p>

<table>
    <thead>
        <tr>
            <th style="width: 22%;">Core Dimension</th>
            <th style="width: 38%;">Primary Objective</th>
            <th style="width: 40%;">Business & Engineering Impact</th>
        </tr>
    </thead>
    <tbody>
        <tr><td><strong>Structural Auditing</strong></td><td>Assess dimensions, feature types, memory footprint, and data schemas.</td><td>Prevents runtime schema errors and optimizes distributed memory allocation.</td></tr>
        <tr><td><strong>Data Quality Assessment</strong></td><td>Identify missingness patterns (MCAR/MAR/MNAR), anomalies, and corruption.</td><td>Determines valid imputation strategies and eliminates corrupted records.</td></tr>
        <tr><td><strong>Distributional Profiling</strong></td><td>Quantify central tendency, dispersion, skewness, and modality.</td><td>Informs necessary non-linear transformations (Log, Box-Cox, Scaling).</td></tr>
        <tr><td><strong>Hypothesis & Pattern Mining</strong></td><td>Uncover non-obvious correlations, multi-way interactions, and clusters.</td><td>Guides high-yield domain feature engineering and hypothesis testing.</td></tr>
        <tr><td><strong>Assumptions Validation</strong></td><td>Test statistical assumptions (normality, homoscedasticity, linearity).</td><td>Ensures correct model selection (Parametric vs. Non-Parametric).</td></tr>
    </tbody>
</table>

<div class="callout">
    <div class="callout-title">📌 CORE PRINCIPLES OF ROBUST EDA</div>
    <ul>
        <li><strong>EDA is iterative and non-linear:</strong> Findings during bivariate analysis often necessitate returning to data cleaning or re-segmenting cohorts.</li>
        <li><strong>Visual diagnostics must always be paired with quantitative statistics:</strong> A chart without skewness, kurtosis, or p-values is prone to subjective human misinterpretation.</li>
        <li><strong>Documenting the EDA journey:</strong> A standardized EDA blueprint serves as the technical single-source-of-truth across cross-functional engineering teams.</li>
    </ul>
</div>

<h1>2. Data Taxonomy & Exploration Strategies</h1>
<table>
    <thead>
        <tr>
            <th style="width: 20%;">Data Type</th>
            <th style="width: 28%;">Definition & Properties</th>
            <th style="width: 26%;">Mathematical Operations</th>
            <th style="width: 26%;">Recommended Visualizations</th>
        </tr>
    </thead>
    <tbody>
        <tr><td><strong>Nominal (Categorical)</strong></td><td>Unordered discrete classes (e.g., Gender, State, Device Type).</td><td>Equality (==, !=), Mode, Frequency Counts.</td><td>Bar charts, Donut charts, Count plots.</td></tr>
        <tr><td><strong>Ordinal (Categorical)</strong></td><td>Ordered discrete categories with ranking (e.g., Education, Rating).</td><td>Comparison (&lt;, &gt;), Median, Percentiles, Rank correlation.</td><td>Ordered Bar charts, Spine plots, Heatmaps.</td></tr>
        <tr><td><strong>Discrete Numeric</strong></td><td>Countable integer quantities (e.g., Number of Dependents, Visits).</td><td>Mean, Median, Standard Deviation, Poisson modeling.</td><td>Step Histograms, Count plots, Stem-and-leaf.</td></tr>
        <tr><td><strong>Continuous Numeric</strong></td><td>Unbroken scale of measurements (e.g., Revenue, Tenure, Speed).</td><td>Full arithmetic, Variance, Kurtosis, Calculus transformations.</td><td>Histograms + KDE, Box plots, ECDF, Violin plots.</td></tr>
        <tr><td><strong>Temporal / Datetime</strong></td><td>Time-series timestamps, cycles (e.g., Transaction Date, Hour).</td><td>Lags, Differences, Seasonality decomposition, Fourier.</td><td>Line charts, Horizon plots, Calendar heatmaps.</td></tr>
    </tbody>
</table>

<h1>3. Data Quality, Missing Values & Outlier Detection</h1>
<h2>3.1 Theoretical Taxonomy of Missing Data Mechanisms</h2>
<ul>
    <li><strong>Missing Completely at Random (MCAR):</strong> Independent of observed/unobserved data. Safe for median/mean imputation.</li>
    <li><strong>Missing at Random (MAR):</strong> Depends systematically on observed features. Requires conditional imputation (KNN/MICE).</li>
    <li><strong>Missing Not at Random (MNAR):</strong> Depends directly on the unobserved value itself. Requires indicator flags.</li>
</ul>

<h2>3.2 Outlier Detection Methodologies</h2>
<pre class="code-block"><span class="code-comment"># 1. Tukey's Interquartile Range (IQR) Rule</span>
Q1 = df['feature'].quantile(0.25)
Q3 = df['feature'].quantile(0.75)
IQR = Q3 - Q1
outliers_iqr = df[(df['feature'] &lt; Q1 - 1.5*IQR) | (df['feature'] &gt; Q3 + 1.5*IQR)]

<span class="code-comment"># 2. Standardized Z-Score Method</span>
z_scores = np.abs((df['feature'] - df['feature'].mean()) / df['feature'].std())
outliers_z = df[z_scores &gt; 3.0]

<span class="code-comment"># 3. Isolation Forest</span>
<span class="code-keyword">from</span> sklearn.ensemble <span class="code-keyword">import</span> IsolationForest
iso = IsolationForest(contamination=0.02, random_state=42)
df['anomaly_flag'] = iso.fit_predict(df[numeric_cols])</pre>

<h1>4. The Multi-Tier Systematic Analysis Framework</h1>
<p>
    • <strong>Tier 1 (Univariate)</strong>: Profiles central tendency, dispersion, skewness, and kurtosis.<br>
    • <strong>Tier 2 (Bivariate)</strong>: Tests relationships via Pearson/Spearman, Welch's t-test, and Chi-Square (&chi;&sup2;).<br>
    • <strong>Tier 3 (Multivariate)</strong>: Evaluates multi-feature interactions, collinearity diagnostics (VIF), and correlation heatmaps.
</p>

<h1>5. Visualization Taxonomy & Decision Matrix</h1>
<table>
    <thead>
        <tr>
            <th style="width: 20%;">Analytical Goal</th>
            <th style="width: 20%;">Variable Configuration</th>
            <th style="width: 25%;">Recommended Visual Chart</th>
            <th style="width: 35%;">Best Practice & Insight Generated</th>
        </tr>
    </thead>
    <tbody>
        <tr><td><strong>Distribution Shape</strong></td><td>1 Continuous</td><td>Histogram + KDE Overlay</td><td>Set appropriate bin widths; overlay mean/median reference lines.</td></tr>
        <tr><td><strong>Spread & Outliers</strong></td><td>1 Continuous</td><td>Box Plot / Violin Plot</td><td>Visualizes IQR, 1.5x whisker boundaries, and distribution density.</td></tr>
        <tr><td><strong>Category Frequency</strong></td><td>1 Categorical</td><td>Horizontal / Vertical Bar Chart</td><td>Sort bars by descending frequency; annotate exact % labels.</td></tr>
        <tr><td><strong>Continuous Association</strong></td><td>2 Continuous</td><td>Scatter Plot + Lowess Trend</td><td>Identifies linear/non-linear correlation, clustering, and heteroscedasticity.</td></tr>
        <tr><td><strong>Grouped Comparison</strong></td><td>1 Cat + 1 Continuous</td><td>Split Violin / Box Plot</td><td>Compares median shift, spread, and tail shape across categories.</td></tr>
        <tr><td><strong>Categorical Interaction</strong></td><td>2 Categorical</td><td>Stacked 100% Bar / Mosaic</td><td>Directly compares relative composition and contingent probabilities.</td></tr>
        <tr><td><strong>Multivariate Correlation</strong></td><td>3+ Continuous</td><td>Masked Heatmap with Annotations</td><td>Use diverging colormaps (coolwarm) with values restricted to [-1, 1].</td></tr>
    </tbody>
</table>

<h1>6. Python Tooling & Library Ecosystem</h1>
<table>
    <thead>
        <tr>
            <th style="width: 22%;">Library / Tool</th>
            <th style="width: 35%;">Core Architectural Role</th>
            <th style="width: 43%;">Key Functions & Unique Capabilities</th>
        </tr>
    </thead>
    <tbody>
        <tr><td><strong>Pandas</strong></td><td>Data manipulation, vectorized transformations, and structural profiling.</td><td><code>read_csv()</code>, <code>describe()</code>, <code>info()</code>, <code>groupby()</code>, <code>pivot_table()</code>.</td></tr>
        <tr><td><strong>NumPy</strong></td><td>High-performance N-dimensional array mathematics.</td><td>Vectorized boolean masking, <code>quantile()</code>, <code>log1p()</code>, <code>where()</code>.</td></tr>
        <tr><td><strong>Matplotlib</strong></td><td>Low-level imperative visualization engine and figure architecture.</td><td>Custom figure grids (<code>subplots</code>), axis formatting, DPI control.</td></tr>
        <tr><td><strong>Seaborn</strong></td><td>High-level statistical visualization library built on Matplotlib.</td><td><code>histplot()</code>, <code>kdeplot()</code>, <code>violinplot()</code>, <code>heatmap()</code>.</td></tr>
        <tr><td><strong>Plotly / Dash</strong></td><td>Declarative, interactive WebGL-rendered charts.</td><td>Hover tooltips, dynamic zooming/panning, 3D surface plots.</td></tr>
        <tr><td><strong>SciPy (scipy.stats)</strong></td><td>Rigorous inferential statistics and hypothesis testing.</td><td><code>ttest_ind()</code>, <code>chi2_contingency()</code>, <code>f_oneway()</code>, <code>shapiro()</code>.</td></tr>
        <tr><td><strong>Missingno</strong></td><td>Dedicated visual diagnostics for missing data matrices.</td><td><code>matrix()</code>, <code>bar()</code>, <code>heatmap()</code> to inspect MCAR/MAR patterns.</td></tr>
    </tbody>
</table>

<h1>7. Practical Demonstration: Telco Churn Dataset</h1>
<p>
    Executed automated EDA on 7,043 records from the Telco Churn dataset downloaded via terminal.
</p>

{fig1_html}
<p>
    &bull; <strong>Data Hygiene:</strong> 11 records in 'TotalCharges' had whitespace characters for new customers (tenure = 0). Imputed with median ($1,397.48).<br>
    &bull; <strong>Target Balance:</strong> Churn rate is 26.5% (1,869 churners vs 5,174 retained).
</p>

{fig2_html}
<p>
    &bull; <strong>Tenure:</strong> Bimodal peaks at Month 1 and Month 72 (Mean: 32.4m, Median: 29.0m).<br>
    &bull; <strong>Monthly Charges:</strong> Mean: $64.76, Median: $70.35.<br>
    &bull; <strong>Total Charges:</strong> Skewness of +0.96.
</p>

{fig3_html}
<p>
    &bull; <strong>Contract Structure:</strong> 55.0% Month-to-Month, 24.1% 2-Year, 20.9% 1-Year.<br>
    &bull; <strong>Internet Service:</strong> 44.0% Fiber Optic, 34.4% DSL, 21.6% No internet.
</p>

{fig4_html}
<p>
    &bull; <strong>Welch's Independent T-Test (Tenure across Churn):</strong> <em>t = -34.82, p = 1.20 &times; 10<sup>-232</sup></em>. Churners have substantially lower median tenure (10m vs 38m).
</p>

{fig5_html}
{fig6_html}
<p>
    &bull; <strong>Chi-Square Test (Contract vs Churn):</strong> <em>&chi;&sup2; = 1184.60, p = 5.86 &times; 10<sup>-258</sup></em>. Month-to-Month contracts have a 42.7% churn rate vs 2.8% for 2-Year contracts.<br>
    &bull; <strong>Fiber Optic Risk:</strong> 41.9% churn rate.
</p>

{fig7_html}

<div class="callout">
    <div class="callout-title">📌 ACTIONABLE FEATURE ENGINEERING RECOMMENDATIONS DERIVED FROM EDA</div>
    <ul>
        <li><strong>Feature 1:</strong> Create <code>Tenure_Cohort</code> (&lt;12m, 12-24m, 24-48m, 48m+).</li>
        <li><strong>Feature 2:</strong> Create <code>Charge_Ratio</code> = (TotalCharges / (tenure * MonthlyCharges)).</li>
        <li><strong>Feature 3:</strong> Create <code>HighRisk_Flag</code> = (Month-to-Month + Electronic Check + Fiber Optic).</li>
    </ul>
</div>

<h1>8. Plan for Documenting, Reporting & Stakeholder Presentation</h1>
<table>
    <thead>
        <tr>
            <th style="width: 22%;">Report Section</th>
            <th style="width: 38%;">Target Audience & Focus</th>
            <th style="width: 40%;">Key Deliverables & Artifacts</th>
        </tr>
    </thead>
    <tbody>
        <tr><td><strong>Executive Summary</strong></td><td>C-Suite & Business Leadership: High-level financial impact and strategic recommendations.</td><td>1-page dashboard, churn drivers ranked by business cost, key ROI projections.</td></tr>
        <tr><td><strong>Data Quality & Governance</strong></td><td>Data Engineering & Compliance: Audit trail of missingness, corruption, and schemas.</td><td>Missing data audit tables, schema validation logs, data pipeline patch recommendations.</td></tr>
        <tr><td><strong>Deep Statistical Analysis</strong></td><td>Data Science & ML Engineers: Distribution parameters, correlation matrices, test stats.</td><td>Full Jupyter notebooks, p-value tables, VIF multicollinearity reports, transformation curves.</td></tr>
        <tr><td><strong>Production Roadmap</strong></td><td>Product Managers & Operations: Concrete operational changes and next sprint milestones.</td><td>Feature engineering dictionary, model baseline criteria, tracking metrics (KPI dashboards).</td></tr>
    </tbody>
</table>

</body>
</html>
"""
    output_doc = os.path.join(SCRIPT_DIR, "EDA_and_Visualization_Framework.doc")
    with open(output_doc, "w", encoding="utf-8") as f:
        f.write(html_content)
    size_kb = os.path.getsize(output_doc) / 1024
    print(f"[+] Successfully generated '{output_doc}' ({size_kb:.2f} KB)")
    return output_doc

if __name__ == "__main__":
    generate_docx()
    generate_html_doc()
